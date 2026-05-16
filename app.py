#!/usr/bin/env python3
"""
Skripsi Helper — AI-powered skripsi title generator & research helper
untuk mahasiswa informatika menggunakan Groq API.
"""

import os
import json
import re
from io import BytesIO
import requests
from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

app = Flask(__name__)

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL = "llama-3.3-70b-versatile"

SEMANTIC_SCHOLAR_URL = "https://api.semanticscholar.org/graph/v1/paper/search"


def groq_chat(messages: list, max_tokens: int = 1024) -> str:
    """Call Groq API and return response text."""
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": GROQ_MODEL,
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": 0.8,
    }
    res = requests.post(GROQ_API_URL, headers=headers, json=payload, timeout=30)
    res.raise_for_status()
    return res.json()["choices"][0]["message"]["content"].strip()


def search_semantic_scholar(query: str, limit: int = 5) -> list:
    """Search existing papers on Semantic Scholar."""
    params = {
        "query": query,
        "limit": limit,
        "fields": "title,authors,year,abstract,url",
    }
    try:
        res = requests.get(SEMANTIC_SCHOLAR_URL, params=params, timeout=10)
        res.raise_for_status()
        data = res.json()
        return data.get("data", [])
    except Exception:
        return []


@app.route("/")
def index():
    return render_template("index.html")


def _set_doc_style(document: Document):
    """Apply clean academic formatting to DOCX."""
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ["Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(14 if style_name == "Heading 1" else 12)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def _add_heading(document: Document, text: str, level: int = 1):
    p = document.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_body_paragraph(document: Document, text: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.5
    p.add_run(text or "-")
    return p


def _add_numbered_list(document: Document, items: list):
    for item in items or []:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(str(item))


def _safe_filename(title: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9_-]+", "-", title.lower()).strip("-")[:60]
    return f"bab-1-{slug or 'outline'}.docx"


@app.route("/api/download-outline", methods=["POST"])
def download_outline():
    body = request.get_json() or {}
    judul = body.get("judul", "Judul Skripsi")
    outline = body.get("outline", {})

    document = Document()
    _set_doc_style(document)

    # Mini cover / title block
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OUTLINE BAB I\nPENDAHULUAN")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(judul.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    _add_heading(document, "BAB I", 1)
    _add_heading(document, "PENDAHULUAN", 1)

    _add_heading(document, "1.1 Latar Belakang", 2)
    _add_body_paragraph(document, outline.get("latar_belakang", ""))

    _add_heading(document, "1.2 Rumusan Masalah", 2)
    _add_numbered_list(document, outline.get("rumusan_masalah", []))

    _add_heading(document, "1.3 Tujuan Penelitian", 2)
    _add_numbered_list(document, outline.get("tujuan_penelitian", []))

    _add_heading(document, "1.4 Manfaat Penelitian", 2)
    manfaat = outline.get("manfaat_penelitian", {}) or {}
    _add_body_paragraph(document, f"Manfaat teoritis: {manfaat.get('teoritis', '-')}")
    _add_body_paragraph(document, f"Manfaat praktis: {manfaat.get('praktis', '-')}")

    _add_heading(document, "1.5 Batasan Masalah", 2)
    _add_numbered_list(document, outline.get("batasan_masalah", []))

    _add_heading(document, "1.6 Sistematika Penulisan", 2)
    _add_numbered_list(document, outline.get("sistematika_penulisan", []))

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=_safe_filename(judul),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/api/generate-titles", methods=["POST"])
def generate_titles():
    """Generate skripsi title suggestions."""
    body = request.get_json()
    bidang = body.get("bidang", "")
    keywords = body.get("keywords", "")
    metode = body.get("metode", "")
    level = body.get("level", "S1")

    prompt = f"""Kamu adalah asisten akademik untuk mahasiswa Indonesia dari berbagai jurusan.
Tugasmu adalah menghasilkan 10 judul skripsi yang kreatif, spesifik, dan layak untuk diteliti.

Informasi dari mahasiswa:
- Bidang/topik: {bidang}
- Kata kunci: {keywords}
- Metode yang diminati: {metode if metode else "bebas"}
- Jenjang: {level}

Aturan judul yang baik:
1. Spesifik dan tidak terlalu umum
2. Mengandung objek penelitian yang jelas
3. Mengandung metode atau pendekatan yang sesuai bidang
4. Bahasa Indonesia yang baku
5. Panjang 10-20 kata
6. Sesuaikan dengan bidang ilmu yang diberikan (bisa sains, sosial, hukum, ekonomi, kesehatan, teknik, dll)

Format output (JSON array):
[
  {{
    "judul": "judul skripsi lengkap",
    "metode": "metode utama yang digunakan",
    "bidang": "sub-bidang ilmu",
    "tingkat_kesulitan": "Mudah/Sedang/Sulit",
    "alasan": "kenapa judul ini menarik dan layak diteliti (1-2 kalimat)"
  }}
]

Hanya output JSON, tidak ada teks lain."""

    try:
        result = groq_chat([{"role": "user", "content": prompt}], max_tokens=2048)
        # Extract JSON from response
        start = result.find("[")
        end = result.rfind("]") + 1
        if start == -1 or end == 0:
            return jsonify({"success": False, "message": "Format response tidak valid"}), 500
        titles = json.loads(result[start:end])
        return jsonify({"success": True, "titles": titles})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500


@app.route("/api/check-research", methods=["POST"])
def check_research():
    """Check existing research on Semantic Scholar."""
    body = request.get_json()
    judul = body.get("judul", "")
    papers = search_semantic_scholar(judul, limit=5)
    return jsonify({"success": True, "papers": papers})


@app.route("/api/suggest-methodology", methods=["POST"])
def suggest_methodology():
    """Suggest research methodology based on title."""
    body = request.get_json()
    judul = body.get("judul", "")

    prompt = f"""Kamu adalah dosen pembimbing skripsi di Indonesia untuk berbagai jurusan.
Berikan saran metodologi penelitian yang detail untuk judul skripsi berikut:

Judul: "{judul}"

Berikan output dalam format JSON:
{{
  "jenis_penelitian": "...",
  "metodologi_utama": "...",
  "langkah_penelitian": ["langkah 1", "langkah 2", "langkah 3", "langkah 4", "langkah 5"],
  "tools_teknologi": ["tool1", "tool2", "tool3"],
  "dataset_saran": "...",
  "estimasi_waktu": "... bulan",
  "referensi_metode": ["paper/buku referensi 1", "paper/buku referensi 2"],
  "tips": "saran singkat untuk mahasiswa"
}}

Hanya output JSON, tidak ada teks lain."""

    try:
        result = groq_chat([{"role": "user", "content": prompt}], max_tokens=1024)
        start = result.find("{")
        end = result.rfind("}") + 1
        methodology = json.loads(result[start:end])
        return jsonify({"success": True, "methodology": methodology})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500


@app.route("/api/generate-outline", methods=["POST"])
def generate_outline():
    """Generate BAB 1 outline based on title."""
    body = request.get_json()
    judul = body.get("judul", "")

    prompt = f"""Kamu adalah dosen pembimbing skripsi di Indonesia untuk berbagai jurusan.
Buatkan outline BAB 1 (Pendahuluan) yang lengkap untuk skripsi dengan judul:

"{judul}"

Format output JSON:
{{
  "latar_belakang": "paragraf latar belakang masalah (3-4 kalimat)",
  "rumusan_masalah": ["rumusan masalah 1", "rumusan masalah 2", "rumusan masalah 3"],
  "tujuan_penelitian": ["tujuan 1", "tujuan 2", "tujuan 3"],
  "manfaat_penelitian": {{
    "teoritis": "manfaat teoritis",
    "praktis": "manfaat praktis"
  }},
  "batasan_masalah": ["batasan 1", "batasan 2", "batasan 3"],
  "sistematika_penulisan": ["BAB I: ...", "BAB II: ...", "BAB III: ...", "BAB IV: ...", "BAB V: ..."]
}}

Hanya output JSON, tidak ada teks lain."""

    try:
        result = groq_chat([{"role": "user", "content": prompt}], max_tokens=1500)
        start = result.find("{")
        end = result.rfind("}") + 1
        outline = json.loads(result[start:end])
        return jsonify({"success": True, "outline": outline})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500


if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
